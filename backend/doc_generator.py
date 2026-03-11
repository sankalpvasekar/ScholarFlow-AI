"""
doc_generator.py — ScholarFlow AI
Generates structured Word documents from Markdown content.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

def generate_docx(markdown_content: str, title: str = "Research Paper") -> io.BytesIO:
    doc = Document()
    
    # ── Document Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Simple Title 
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ── Basic Markdown Parser ──
    lines = markdown_content.split('\n')
    
    for line in lines:
        raw_line = line.strip()
        if not raw_line:
            continue
            
        # Headers
        if raw_line.startswith('### '):
            doc.add_heading(raw_line[4:], level=3)
        elif raw_line.startswith('## '):
            doc.add_heading(raw_line[3:], level=2)
        elif raw_line.startswith('# '):
            # If it's the main title (first #), we already handled it or we can skip heading 0
            if doc.paragraphs and doc.paragraphs[0].text == title:
                continue
            doc.add_heading(raw_line[2:], level=1)
            
        # Lists
        elif raw_line.startswith('* ') or raw_line.startswith('- ') or re.match(r'^\d+\.', raw_line):
            text = re.sub(r'^(\*|-|\d+\.)\s+', '', raw_line)
            doc.add_paragraph(text, style='List Bullet' if not raw_line[0].isdigit() else 'List Number')
        
        # Paragraphs
        else:
            para = doc.add_paragraph()
            # Handle bold **text** and italic *text*
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', raw_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target
